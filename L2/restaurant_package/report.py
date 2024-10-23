# restaurant_package/report.py
from docx import Document
from restaurant_package.order import Order
from restaurant_package.table import Table


def save_to_doc(order: Order, table: Table, filename: str):
    doc = Document()

    # Заголовок
    doc.add_heading('Отчет по заказу в ресторане', 0)

    # Столик
    doc.add_heading('Информация о столике', level=1)
    doc.add_paragraph(f"Номер столика: {table.table_number}")
    doc.add_paragraph(f"Количество мест: {table.seats}")
    doc.add_paragraph(f"Статус: {'Занят' if table.is_occupied else 'Свободен'}")

    # Заказ
    doc.add_heading('Заказ', level=1)
    for dish in order.dishes:
        doc.add_paragraph(f"{dish.name}: {dish.price:.2f} руб.")

    # Итоговая сумма и скидка
    doc.add_paragraph(f"\nИтоговая сумма с учетом скидки ({order.discount}%): {order.calculate_total():.2f} руб.")

    # Сохранение документа
    doc.save(filename)
